import pdfplumber
from docx import Document
from pathlib import Path


class ResumeParser:
    """Extracts clean text from resume files (PDF, DOCX, TXT)."""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt"}

    def extract(self, file_path: str) -> dict:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format '{ext}'. "
                f"Supported: {self.SUPPORTED_FORMATS}"
            )

        if ext == ".pdf":
            raw_text = self._extract_pdf(path)
        elif ext == ".docx":
            raw_text = self._extract_docx(path)
        else:
            raw_text = path.read_text(encoding="utf-8")

        text = self._clean_text(raw_text)

        return {
            "file_name": path.name,
            "format": ext,
            "text": text,
            "word_count": len(text.split()),
            "char_count": len(text),
        }

    def _extract_pdf(self, path: Path) -> str:
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n".join(pages)

    def _extract_docx(self, path: Path) -> str:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    def _clean_text(self, text: str) -> str:
        lines = [line.strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line)